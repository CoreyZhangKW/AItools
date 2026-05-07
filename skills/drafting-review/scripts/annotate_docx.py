#!/usr/bin/env python3
"""
Annotate a Word (.docx) file using issues from issues.json.

This script is intentionally dependency-light and robust:
- It does NOT rely on Microsoft Word APIs.
- It does NOT attempt to create native Word "Comments" (which python-docx doesn't support reliably).
- Instead, it appends clearly-formatted inline annotations near the relevant text.

Usage:
  python annotate_docx.py <input.docx> <issues.json> <output_reviewed.docx>
"""

from __future__ import annotations

import json
import os
import sys
from dataclasses import dataclass
from typing import Any, Dict, Iterable, List, Optional, Tuple


try:
    from docx import Document  # type: ignore
    from docx.shared import Pt, RGBColor  # type: ignore
except Exception as e:  # pragma: no cover
    print(
        "[ERROR] Missing dependency: python-docx. Install with:\n"
        "  python -m pip install python-docx\n"
        f"Details: {e}"
    )
    sys.exit(1)


@dataclass(frozen=True)
class Issue:
    location: str
    issue_type: str
    original_text: str
    suggestion: str
    reason: str


def _load_issues(path: str) -> List[Issue]:
    if not os.path.exists(path):
        print(f"[ERROR] issues.json not found: {path}")
        sys.exit(1)
    try:
        with open(path, "r", encoding="utf-8") as f:
            raw = json.load(f)
    except Exception as e:
        print(f"[ERROR] Failed to parse JSON: {path}\n{e}")
        sys.exit(1)

    if not isinstance(raw, list):
        print("[ERROR] issues.json must be a JSON array.")
        sys.exit(1)

    issues: List[Issue] = []
    for i, item in enumerate(raw):
        if not isinstance(item, dict):
            print(f"[ERROR] issues.json entry #{i} is not an object.")
            sys.exit(1)
        try:
            issues.append(
                Issue(
                    location=str(item.get("location", "")).strip(),
                    issue_type=str(item.get("issue_type", "")).strip(),
                    original_text=str(item.get("original_text", "")).strip(),
                    suggestion=str(item.get("suggestion", "")).strip(),
                    reason=str(item.get("reason", "")).strip(),
                )
            )
        except Exception as e:
            print(f"[ERROR] issues.json entry #{i} invalid: {e}")
            sys.exit(1)

    # Keep only issues with something actionable.
    issues = [x for x in issues if x.original_text or x.location or x.suggestion or x.reason]
    return issues


def _iter_all_paragraphs(doc: Document) -> Iterable[Any]:
    # Body paragraphs
    for p in doc.paragraphs:
        yield p
    # Tables (including nested)
    stack = list(doc.tables)
    while stack:
        t = stack.pop()
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                # Nested tables in cells
                for nt in cell.tables:
                    stack.append(nt)


def _normalize(s: str) -> str:
    return " ".join((s or "").replace("\u00A0", " ").split())


def _format_annotation(issue: Issue) -> str:
    # Keep it compact and scan-friendly.
    parts: List[str] = []
    if issue.issue_type:
        parts.append(f"【{issue.issue_type}】")
    if issue.suggestion:
        parts.append(f"建议：{issue.suggestion}")
    if issue.reason:
        parts.append(f"原因：{issue.reason}")
    if issue.location:
        parts.append(f"位置：{issue.location}")
    return " ".join(parts).strip()


def _append_annotation(paragraph: Any, note: str) -> None:
    # Visual style: small red italic text, prefixed with a tag.
    run = paragraph.add_run(f"\n[REVIEW] {note}")
    try:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0xB0, 0x00, 0x20)  # deep red
    except Exception:
        # If styles fail, we still keep the text.
        pass


def annotate_docx(input_docx: str, issues_json: str, output_docx: str) -> Tuple[int, int]:
    if not os.path.exists(input_docx):
        print(f"[ERROR] Input .docx not found: {input_docx}")
        sys.exit(1)

    issues = _load_issues(issues_json)
    doc = Document(input_docx)

    located = 0
    unlocated: List[Issue] = []

    # Build paragraph index once for speed.
    paragraphs = list(_iter_all_paragraphs(doc))
    paragraph_text_norm = [_normalize(p.text) for p in paragraphs]

    for issue in issues:
        needle = _normalize(issue.original_text)
        if needle:
            # Find first matching paragraph. (Deliberate: stable + predictable output.)
            idx = next((i for i, t in enumerate(paragraph_text_norm) if needle in t), None)
        else:
            idx = None

        if idx is None:
            unlocated.append(issue)
            continue

        _append_annotation(paragraphs[idx], _format_annotation(issue))
        located += 1

    if unlocated:
        doc.add_page_break()
        p = doc.add_paragraph("Unlocated issues (text not found; review manually):")
        try:
            p.runs[0].font.bold = True
        except Exception:
            pass
        for issue in unlocated:
            doc.add_paragraph(f"- {_format_annotation(issue)}")

    doc.save(output_docx)
    return located, len(unlocated)


if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python annotate_docx.py <input.docx> <issues.json> <output_reviewed.docx>")
        sys.exit(1)

    located, unlocated = annotate_docx(sys.argv[1], sys.argv[2], sys.argv[3])
    print(f"[OK] Saved: {sys.argv[3]}")
    print(f"[OK] Issues annotated: {located}")
    if unlocated:
        print(f"[WARN] Issues not located in text: {unlocated} (listed at end of document)")

