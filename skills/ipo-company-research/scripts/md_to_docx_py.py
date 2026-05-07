#!/usr/bin/env python3
"""Convert Markdown to Word (.docx) using python-docx."""
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

input_path = sys.argv[1]
output_path = sys.argv[2]

with open(input_path, 'r', encoding='utf-8') as f:
    content = f.read()

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(13)

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def process_text(paragraph, text):
    paragraph.clear()
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)
        for r in paragraph.runs:
            r.font.name = 'SimSun'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_table(doc, header, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(header))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(header):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        run = p.add_run(col)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1A3C6E')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = table.rows[r+1].cells[j]
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9)
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if r % 2 == 0:
                set_cell_shading(cell, 'F2F6FC')
    doc.add_paragraph()

lines = content.split('\n')
i = 0
in_table = False
table_header = []
table_rows = []

while i < len(lines):
    line = lines[i]
    
    if line.strip().startswith('```'):
        i += 1
        continue
    
    if line.strip() == '---':
        i += 1
        continue
    
    if line.startswith('# ') and not line.startswith('## '):
        doc.add_heading(line[2:], level=1)
        i += 1
        continue
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
        i += 1
        continue
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
        i += 1
        continue
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
        i += 1
        continue
    
    # Table
    if '|' in line and line.strip().startswith('|'):
        if not in_table:
            in_table = True
            table_header = [c.strip() for c in line.split('|')[1:-1]]
            i += 1
            if i < len(lines) and re.match(r'^[\|\s\-:]+$', lines[i].strip()):
                i += 1
        else:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                table_rows.append(cells)
        i += 1
        if i >= len(lines) or '|' not in lines[i]:
            if table_header and table_rows:
                add_table(doc, table_header, table_rows)
            in_table = False
            table_header = []
            table_rows = []
        continue
    
    if not line.strip():
        i += 1
        continue
    
    # Ordered/unordered list
    if line.strip().startswith('- ') or re.match(r'^\d+\.\s', line.strip()):
        p = doc.add_paragraph()
        process_text(p, line.strip())
        p.paragraph_format.line_spacing = Pt(18)
        i += 1
        continue
    
    p = doc.add_paragraph()
    process_text(p, line)
    p.paragraph_format.line_spacing = Pt(20)
    i += 1

if in_table and table_header and table_rows:
    add_table(doc, table_header, table_rows)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

doc.save(output_path)
print(f"[OK] Saved: {output_path}")
