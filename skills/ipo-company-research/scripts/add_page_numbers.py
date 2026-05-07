#!/usr/bin/env python3
"""Add page numbers to footer of existing docx."""
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

input_path = sys.argv[1]
output_path = input_path  # overwrite

doc = Document(input_path)

for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Clear existing
    p.clear()
    
    # Add "第 " text
    run1 = p.add_run('第 ')
    run1.font.size = Pt(9)
    run1.font.name = 'SimSun'
    
    # Add page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run_page = p.add_run()
    run_page._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run_page._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_page._r.append(fldChar2)
    
    # Add " 页" text
    run2 = p.add_run(' 页')
    run2.font.size = Pt(9)
    run2.font.name = 'SimSun'

doc.save(output_path)
print(f"[OK] Page numbers added: {output_path}")
